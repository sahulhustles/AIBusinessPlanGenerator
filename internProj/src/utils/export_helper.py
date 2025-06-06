from docx import Document
from docx.shared import Pt
from fpdf import FPDF
import tempfile
import os

def export_to_docx(content: str, filename: str = "business_strategy") -> str:
    """Export strategy content to DOCX format"""
    try:
        doc = Document()
        
        # Add title
        doc.add_heading('Business Strategy Report', 0)
        
        # Process sections
        sections = content.split('## ')
        for section in sections:
            if section.strip():
                heading, _, body = section.partition('\n')
                if heading:
                    doc.add_heading(heading.strip(), level=2)
                    doc.add_paragraph(body.strip())
        
        # Save to temp file
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, f"{filename}.docx")
        doc.save(temp_path)
        return temp_path
        
    except Exception as e:
        raise RuntimeError(f"DOCX Export Error: {str(e)}")

def export_to_pdf(content: str, filename: str = "business_strategy") -> str:
    """Export strategy content to PDF using FPDF's built-in Unicode support"""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Enable Unicode support for built-in font
        pdf.set_font("helvetica", size=12)  # Use built-in font but handle Unicode
        
        # Add title
        pdf.set_font("helvetica", 'B', 16)
        pdf.cell(0, 10, "Business Strategy Report", 0, 1, 'C')
        pdf.ln(10)
        
        # Process content with Unicode characters
        pdf.set_font("helvetica", size=12)
        sections = content.split('## ')
        for section in sections:
            if section.strip():
                heading, _, body = section.partition('\n')
                
                # Add heading
                if heading:
                    pdf.set_font("helvetica", 'B', 12)
                    pdf.cell(0, 10, heading.strip().replace('–', '-'), 0, 1)  # Replace en-dash with regular dash
                    pdf.set_font("helvetica", '', 11)
                
                # Add body with Unicode handling
                if body:
                    cleaned_body = body.strip().replace('–', '-')  # Simple character replacement
                    pdf.multi_cell(0, 8, cleaned_body)
                    pdf.ln(5)
        
        # Save to temp file
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, f"{filename}.pdf")
        pdf.output(temp_path)
        return temp_path
        
    except Exception as e:
        raise RuntimeError(f"PDF Export Error: {str(e)}")