from docx import Document
from docx.shared import Pt
import pdfkit
import os
import tempfile

def export_to_docx(content: str, filename: str = "business_strategy") -> str:
    """Export strategy content to DOCX format"""
    try:
        doc = Document()
        
        # Add title
        doc.add_heading('Business Strategy Report', 0)
        
        # Process sections
        sections = content.split('## ')
        for section in sections:
            if section.strip():
                heading, _, body = section.partition('\n')
                if heading:
                    doc.add_heading(heading.strip(), level=2)
                    doc.add_paragraph(body.strip())
        
        # Save to temp file
        temp_path = os.path.join(tempfile.gettempdir(), f"{filename}.docx")
        doc.save(temp_path)
        return temp_path
        
    except Exception as e:
        raise RuntimeError(f"DOCX Export Error: {str(e)}")

def export_to_pdf(content: str, filename: str = "business_strategy") -> str:
    """Export strategy content to PDF format"""
    try:
        # Convert markdown to HTML
        html_content = f"""
        <html>
            <head>
                <style>
                    body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
                    h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; }}
                    h2 {{ color: #3498db; margin-top: 1.5em; }}
                </style>
            </head>
            <body>
                <h1>Business Strategy Report</h1>
                {content.replace('## ', '<h2>').replace('\n', '<br>')}
            </body>
        </html>
        """
        
        # PDF options
        options = {
            'encoding': 'UTF-8',
            'quiet': '',
            'margin-top': '0.5in',
            'margin-right': '0.5in',
            'margin-bottom': '0.5in',
            'margin-left': '0.5in'
        }
        
        # Save to temp file
        temp_path = os.path.join(tempfile.gettempdir(), f"{filename}.pdf")
        pdfkit.from_string(html_content, temp_path, options=options)
        return temp_path
        
    except Exception as e:
        raise RuntimeError(f"PDF Export Error: {str(e)}")