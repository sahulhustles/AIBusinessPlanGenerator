from docx import Document
from docx.shared import Pt
from fpdf import FPDF
import tempfile
import os

def export_to_docx(content: str, filename: str = "business_strategy") -> str:
    """Export strategy content to DOCX format"""
    try:
        doc = Document()
        
        # Add title
        doc.add_heading('Business Strategy Report', 0)
        
        # Process sections
        sections = content.split('## ')
        for section in sections:
            if section.strip():
                heading, _, body = section.partition('\n')
                if heading:
                    doc.add_heading(heading.strip(), level=2)
                    doc.add_paragraph(body.strip())
        
        # Save to temp file
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, f"{filename}.docx")
        doc.save(temp_path)
        return temp_path
        
    except Exception as e:
        raise RuntimeError(f"DOCX Export Error: {str(e)}")

def export_to_pdf(content: str, filename: str = "business_strategy") -> str:
    """Export strategy content to PDF using FPDF"""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Set fonts
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, "Business Strategy Report", 0, 1, 'C')
        pdf.ln(10)
        
        # Process content
        sections = content.split('## ')
        for section in sections:
            if section.strip():
                heading, _, body = section.partition('\n')
                
                # Add heading
                if heading:
                    pdf.set_font("Arial", 'B', 12)
                    pdf.cell(0, 10, heading.strip(), 0, 1)
                
                # Add body
                if body:
                    pdf.set_font("Arial", size=11)
                    pdf.multi_cell(0, 8, body.strip())
                    pdf.ln(5)
        
        # Save to temp file
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, f"{filename}.pdf")
        pdf.output(temp_path)
        return temp_path
        
    except Exception as e:
        raise RuntimeError(f"PDF Export Error (FPDF): {str(e)}")